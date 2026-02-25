import json
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from lxml import etree
from matplotlib import font_manager
from typing import List, Optional

class DocxGenerator: 
	def __init__(self, docx_path: Optional[str] = None): 
		if docx_path: 
			self.docx = Document(docx_path)
		else: 
			self.docx = Document()
		self._set_margin()
		self._set_font()
		self.paragraphs: List[Paragraph] = []
		self.current_paragraph_index: int = -1
		

	def _set_margin(self, 
			top: float = 1.5, bottom: float = 1.5, 
			left: float = 1.5, right: float = 1.5
		) -> None:
		section = self.docx.sections[0]
		section.top_margin = Cm(top)
		section.bottom_margin = Cm(bottom)
		section.left_margin = Cm(left)
		section.right_margin = Cm(right)

	def _set_font(self, 
			font_name: str = "Hiragino Mincho ProN", 
			font_size: int = 10
		) -> None:
		self.font = font_name
		self.size = font_size
		style = self.docx.styles["Normal"]
		assert isinstance(style, ParagraphStyle)
		style.font.name = font_name
		style.font.size = Pt(font_size)

		r_pr = style.element.rPr # type: ignore[attr-defined]
		r_fonts = r_pr.rFonts # type: ignore[attr-defined]
		
		r_fonts.set(qn("w:eastAsia"), font_name)
		r_fonts.set(qn("w:ascii"), font_name)
		r_fonts.set(qn("w:hAnsi"), font_name)
		r_fonts.set(qn("w:cs"), font_name)

	def add_paragraph(self, text: str = "", space_after: float = 0.0) -> None: 
		paragraph = self.docx.add_paragraph(text)
		paragraph.paragraph_format.space_after = Cm(space_after)
		self.paragraphs.append(paragraph)
		self.current_paragraph_index += 1

	def add_run(self, text: str = "") -> None: 
		if self.current_paragraph_index == -1: 
			self.add_paragraph()
		paragraph = self.paragraphs[self.current_paragraph_index]
		paragraph.add_run(text)

	def _add_run_properties(self, 
			run_elem: etree._Element, font_name: str, 
			lang_east_asia: str, font_size: Optional[int] = None
		) -> None: 
		r_pr = etree.SubElement(run_elem, qn("w:rPr"))

		r_fonts = etree.SubElement(r_pr, qn("w:rFonts"))
		r_fonts.set(qn("w:hint"), "eastAsia")
		r_fonts.set(qn("w:ascii"), font_name)
		r_fonts.set(qn("w:hAnsi"), font_name)
		r_fonts.set(qn("w:eastAsia"), font_name)
		r_fonts.set(qn("w:cs"), font_name)

		if font_size is not None: 
			sz = etree.SubElement(r_pr, qn("w:sz"))
			sz.set(qn("w:val"), str(font_size * 2))
			sz_cs = etree.SubElement(r_pr, qn("w:szCs"))
			sz_cs.set(qn("w:val"), str(font_size * 2))

		lang = etree.SubElement(r_pr, qn("w:lang"))
		lang.set(qn("w:val"), "en-US")
		lang.set(qn("w:eastAsia"), lang_east_asia)

	def add_field0(self, text: str, notation: str, align_type: int) -> None: 
		if self.current_paragraph_index == -1: 
			self.add_paragraph()
		p_elem = self.paragraphs[self.current_paragraph_index]._element
		runs = []
		
		runs = []
		r1 = etree.Element(qn("w:r"))
		self._add_run_properties(r1, self.font, "ja-JP")
		fld_begin = etree.SubElement(r1, qn("w:fldChar"))
		fld_begin.set(qn("w:fldCharType"), "begin")
		runs.append(r1)

		r2 = etree.Element(qn("w:r"))
		self._add_run_properties(r2, self.font, "ja-JP")
		instr = etree.SubElement(r2, qn("w:instrText"))
		instr.set(qn("xml:space"), "preserve")
		instr.text = (
			" EQ \\* jc{} \\* \"{}\" \\* hps{} \\o \\ad(\\s \\up {}({}),{})"
		).format(align_type, self.font, self.size, self.size, notation, text)
		runs.append(r2)

		r3 = etree.Element(qn("w:r"))
		self._add_run_properties(r3, self.font, "ja-JP")
		fld_end = etree.SubElement(r3, qn("w:fldChar"))
		fld_end.set(qn("w:fldCharType"), "end")
		runs.append(r3)

		for r in runs: 
			p_elem.append(r)

	def add_field1(self, text: str, notation_up: str, notation_down: str) -> None: 
		if self.current_paragraph_index == -1: 
			self.add_paragraph()
		self.add_run(text)
		p_elem = self.paragraphs[self.current_paragraph_index]._element
		runs = []
		
		runs = []
		r1 = etree.Element(qn("w:r"))
		self._add_run_properties(r1, self.font, "ja-JP")
		fld_begin = etree.SubElement(r1, qn("w:fldChar"))
		fld_begin.set(qn("w:fldCharType"), "begin")
		runs.append(r1)

		r2 = etree.Element(qn("w:r"))
		self._add_run_properties(r2, self.font, "ja-JP", font_size=(self.size // 2))
		instr = etree.SubElement(r2, qn("w:instrText"))
		instr.set(qn("xml:space"), "preserve")
		instr.text = (
			" EQ \\o \\al (\\s \\up {}({}),\\s \\do 0 ({}))"
		).format(self.size // 2, notation_up, notation_down)
		runs.append(r2)

		r3 = etree.Element(qn("w:r"))
		self._add_run_properties(r3, self.font, "ja-JP")
		fld_end = etree.SubElement(r3, qn("w:fldChar"))
		fld_end.set(qn("w:fldCharType"), "end")
		runs.append(r3)

		for r in runs: 
			p_elem.append(r)

	def save(self, path: str) -> None: 
		self.docx.save(path)

	def set_columns(self, num_columns: int, space: int = 425) -> None: 
		section = self.docx.sections[0]
		sect_pr = section._sectPr

		cols = sect_pr.find(qn("w:cols"))
		if cols is None: 
			cols = OxmlElement("w:cols")
			sect_pr.append(cols) # type: ignore[union-attr]

		cols.set(qn("w:num"), str(num_columns))
		cols.set(qn("w:space"), str(space))

if __name__ == "__main__": 

	DIRPATH = os.path.dirname(os.path.abspath(__file__))
	SRCPATH = os.path.join(DIRPATH, "..", "src")